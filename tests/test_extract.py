"""Tests for scripts/extract.py — TXT and DOCX extraction."""

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent / 'scripts'))

from extract import DocumentExtractor


class TestExtractTxt:
    def test_basic_txt_extraction(self, tmp_path):
        # Create a sample txt file
        txt_file = tmp_path / "raw" / "test_sop.txt"
        txt_file.parent.mkdir(parents=True)
        txt_file.write_text("This is a test document.\nIt has two lines.", encoding='utf-8')

        processed_dir = tmp_path / "processed"
        processed_dir.mkdir()

        extractor = DocumentExtractor(
            raw_dir=str(txt_file.parent),
            processed_dir=str(processed_dir)
        )
        chunks = extractor.extract_txt(txt_file)

        assert len(chunks) == 1
        assert chunks[0]['doc_name'] == 'test_sop'
        assert chunks[0]['doc_type'] == 'txt'
        assert 'two lines' in chunks[0]['text']

    def test_empty_txt(self, tmp_path):
        txt_file = tmp_path / "empty.txt"
        txt_file.write_text("", encoding='utf-8')

        extractor = DocumentExtractor(
            raw_dir=str(tmp_path),
            processed_dir=str(tmp_path)
        )
        chunks = extractor.extract_txt(txt_file)
        # Empty file produces no meaningful chunks
        assert len(chunks) == 0 or chunks[0]['text'].strip() == ''


class TestExtractDocxVirtualPages:
    def test_virtual_page_numbering(self, tmp_path):
        """Verify DOCX extraction creates multiple virtual pages."""
        try:
            from docx import Document
        except ImportError:
            import pytest
            pytest.skip("python-docx not installed")

        # Create a DOCX with enough text to span multiple virtual pages
        doc = Document()
        for i in range(100):
            doc.add_paragraph(f"This is paragraph {i} with some filler text to increase character count. " * 3)

        docx_file = tmp_path / "test.docx"
        doc.save(str(docx_file))

        extractor = DocumentExtractor(
            raw_dir=str(tmp_path),
            processed_dir=str(tmp_path)
        )
        chunks = extractor.extract_docx(docx_file, chars_per_page=3000)

        # Should produce more than 1 virtual page
        assert len(chunks) > 1
        # Pages should be sequential
        pages = [c['page'] for c in chunks]
        assert pages == list(range(1, len(pages) + 1))

    def test_short_docx_single_page(self, tmp_path):
        """A short DOCX should produce a single virtual page."""
        try:
            from docx import Document
        except ImportError:
            import pytest
            pytest.skip("python-docx not installed")

        doc = Document()
        doc.add_paragraph("Short document.")
        docx_file = tmp_path / "short.docx"
        doc.save(str(docx_file))

        extractor = DocumentExtractor(
            raw_dir=str(tmp_path),
            processed_dir=str(tmp_path)
        )
        chunks = extractor.extract_docx(docx_file)

        assert len(chunks) == 1
        assert chunks[0]['page'] == 1
